from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = PROJECT_ROOT / "report_assets"
SUMMARY_PATH = ASSET_DIR / "experiment_summary.json"
OUTPUT_PATH = PROJECT_ROOT / "PCB_defect_YOLO_course_report_expanded.docx"

CLASSES_ZH = {
    "missing_pad": "焊盘缺失",
    "mouse_bite": "鼠咬",
    "open_circuit": "开路",
    "short": "短路",
    "spur": "毛刺",
    "spurious_copper": "残铜",
}


def set_run_font(run, size: float | None = None, bold: bool | None = None, name: str = "宋体") -> None:
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line: bool = True, align=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.first_line_indent = Cm(0.74) if first_line else None


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    for style_name, size, font, bold in [
        ("Heading 1", 16, "黑体", True),
        ("Heading 2", 14, "宋体", True),
        ("Heading 3", 12, "黑体", True),
    ]:
        try:
            style = doc.styles[style_name]
        except KeyError:
            style = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), font)
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 20)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 1" else 12)


def clear_document(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def setup_page(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.NEW_PAGE
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)


def add_centered(doc: Document, text: str, size: float = 12, bold: bool = False, font: str = "宋体") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, name=font)


def add_body_paragraph(doc: Document, text: str, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    set_paragraph_format(p, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, size=12, name="宋体")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Paragraph")
        p.paragraph_format.left_indent = Cm(0.75)
        p.paragraph_format.first_line_indent = Cm(-0.35)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run("· " + item)
        set_run_font(run, size=12, name="宋体")


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 14 if level == 2 else 12, bold=True, name="黑体" if level in {1, 3} else "宋体")


def set_cell_text(cell, text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_after = Pt(0)
    p.text = ""
    run = p.add_run(text)
    set_run_font(run, size=10.5, bold=bold, name="宋体")


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "9CA3AF")
        borders.append(tag)
    tbl_pr.append(borders)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], caption: str) -> None:
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(6)
    cap.paragraph_format.space_after = Pt(4)
    run = cap.add_run(caption)
    set_run_font(run, size=10.5, name="宋体")

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    try:
        table.style = "Table Grid"
    except KeyError:
        set_table_borders(table)
    for i, header in enumerate(headers):
        shade_cell(table.rows[0].cells[i], "E5E7EB")
        set_cell_text(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()


def add_picture(doc: Document, image_path: Path, caption: str, width: float = 5.8) -> None:
    if not image_path.exists():
        add_body_paragraph(doc, f"（图片缺失：{image_path.name}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    cap_run = cap.add_run(caption)
    set_run_font(cap_run, size=10.5, name="宋体")


def fmt(value: float) -> str:
    return f"{value:.5f}"


def load_summary() -> dict[str, object]:
    if not SUMMARY_PATH.exists():
        raise FileNotFoundError(f"Missing report asset summary: {SUMMARY_PATH}")
    return json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))


def build_report() -> None:
    templates = sorted(PROJECT_ROOT.glob("*.docx"))
    if not templates:
        raise FileNotFoundError("No DOCX template found in project root.")
    template = next(path for path in templates if path.name != OUTPUT_PATH.name)
    summary = load_summary()
    dataset = summary["dataset"]
    runs = summary["runs"]

    doc = Document(template)
    clear_document(doc)
    setup_page(doc)
    configure_styles(doc)

    add_centered(doc, "计算机视觉课程综合报告", size=22, bold=True, font="黑体")
    doc.add_paragraph()
    add_centered(doc, "基于 YOLO11/YOLO26 的 PCB 缺陷小目标检测实验研究与扩展分析", size=16, bold=True, font="黑体")
    doc.add_paragraph()
    for label, value in [
        ("课程名称", "计算机视觉"),
        ("学院", "信息科学与技术学院"),
        ("专业", "计算机技术 / 软件工程 / 电子信息"),
        ("学号", "待填写"),
        ("姓名", "待填写"),
        ("指导教师", "待填写"),
        ("提交日期", "2026 年 7 月 1 日"),
    ]:
        add_centered(doc, f"{label}：{value}", size=12, font="宋体")
    doc.add_page_break()

    add_heading(doc, "摘要", 1)
    add_body_paragraph(
        doc,
        "PCB 缺陷检测是工业自动光学检测（AOI）中的典型计算机视觉任务。与普通自然场景目标检测相比，PCB 图像具有纹理重复、缺陷尺度小、类别间形态差异细微、背景与目标边界不明显等特点，直接采用全图缩放训练容易造成局部缺陷信息损失。本文围绕公开 PCB 缺陷数据集，完成了从 COCO 标注到 YOLO 格式的转换、训练集/验证集/测试集划分、标签质量检查、YOLO11n 基线训练、输入分辨率扩展实验以及 YOLO26n 对比实验。实验结果表明，在相同 100 epoch 训练设定下，YOLO11n 从 640 输入提升到 960 输入后，mAP@50 从 0.79002 提升到 0.85118，Precision 从 0.79568 提升到 0.89312，但 Recall 只从 0.74765 提升到 0.78091，说明提升输入分辨率能够改善整体检测质量，却仍不能完全解决小缺陷漏检问题。进一步对比 YOLO26n-960 发现，其在本数据集上的最终 mAP@50 为 0.76921，低于 YOLO11n-960，说明通用模型版本升级并不必然带来特定工业小目标场景收益。基于上述结果，本文将课程项目扩展为“数据审计—分辨率消融—模型族对比—误差分析—自适应切片方案设计”的完整实验链路，并提出后续引入局部切片、困难样本挖掘和轻量化部署评估的改进方向。"
    )
    add_body_paragraph(doc, "关键词：PCB 缺陷检测；YOLO11；YOLO26；小目标检测；目标检测；误差分析", first_line=False)
    add_heading(doc, "Abstract", 1)
    add_body_paragraph(
        doc,
        "PCB defect detection is a representative industrial computer vision task. Compared with general object detection in natural scenes, PCB inspection involves repetitive textures, tiny defects, subtle inter-class appearance differences, and weak boundaries between defects and background. This report builds a complete experimental pipeline for PCB defect detection, including COCO-to-YOLO conversion, dataset split and audit, YOLO11n baseline training, input-resolution ablation, YOLO26n comparison, metric visualization, and failure analysis. The results show that increasing the input size of YOLO11n from 640 to 960 improves mAP@50 from 0.79002 to 0.85118 and precision from 0.79568 to 0.89312, while recall improves only moderately from 0.74765 to 0.78091. YOLO26n at 960 does not outperform YOLO11n-960 on this dataset, indicating that task-specific small-object strategies are still necessary. The project is therefore extended from a single baseline into a reproducible course project with data analysis, controlled experiments, qualitative diagnosis, and an adaptive slicing research plan."
    )
    add_body_paragraph(doc, "Key words: PCB defect detection; YOLO11; YOLO26; small object detection; object detection; failure analysis", first_line=False)
    doc.add_page_break()

    add_heading(doc, "目 录", 1)
    for item in [
        "1 绪论",
        "2 计算机视觉基础理论与相关算法",
        "3 数据集构建与实验方案",
        "4 实验结果与扩展分析",
        "5 误差分析与后续改进方案",
        "6 总结与展望",
        "参考文献",
        "附录 A 核心代码与汇报口径",
    ]:
        add_body_paragraph(doc, item, first_line=False)
    doc.add_page_break()

    add_heading(doc, "1 绪论", 1)
    add_heading(doc, "1.1 研究背景与意义", 2)
    add_body_paragraph(
        doc,
        "印刷电路板（Printed Circuit Board, PCB）是电子设备中承载元器件连接关系的基础部件。随着电子产品向高密度、小型化和高可靠性方向发展，PCB 生产线对缺陷检测的实时性与准确率提出了更高要求。传统人工目检受经验、疲劳和主观判断影响较大，难以满足批量化生产中稳定、连续、可追溯的质量控制需求。自动光学检测系统能够利用图像采集和视觉算法对焊盘缺失、开路、短路、毛刺、鼠咬、残铜等缺陷进行定位与分类，是计算机视觉在智能制造领域的重要应用。"
    )
    add_body_paragraph(
        doc,
        "本课题以 PCB 缺陷目标检测为研究对象，选择 YOLO 系列轻量化检测模型作为主要实验框架。YOLO 模型具有端到端训练、推理速度快、工程部署生态完善等优势，适合课程项目在有限时间内完成从数据准备到模型训练、评估和可视化分析的完整链路。然而，PCB 缺陷往往占整幅图像面积比例很小，全图缩放到固定输入尺寸后，缺陷边缘、局部纹理和细小形态容易被压缩，导致漏检或定位不稳定。因此，仅完成一个 baseline 还不足以支撑完整课程报告，需要通过实验对照和误差分析扩展出明确的问题意识。"
    )
    add_heading(doc, "1.2 国内外研究现状", 2)
    add_body_paragraph(
        doc,
        "工业视觉检测早期主要依赖阈值分割、边缘检测、形态学处理、模板匹配与手工特征分类。此类方法在照明受控、背景简单、缺陷形态稳定的场景下具有可解释性强、计算量低的优势，但在复杂纹理、批次差异和微小缺陷条件下泛化能力有限。深度学习方法通过卷积神经网络自动学习多层次特征，大幅提升了目标检测、分类和分割任务的鲁棒性。两阶段检测器如 Faster R-CNN 具有较高定位精度，但推理成本较高；一阶段检测器如 SSD、YOLO 系列更强调速度与精度平衡，适合工业在线检测。"
    )
    add_body_paragraph(
        doc,
        "近年来，检测模型的发展呈现两条趋势：一是持续改进卷积特征提取、多尺度融合、标签分配和训练策略；二是引入 Transformer 或端到端检测思想以减少后处理依赖。Ultralytics 文档显示，YOLO11 在模型结构和训练方法上相对前代进行了效率与精度优化；YOLO26 则进一步强调原生端到端推理、轻量检测头和针对小目标覆盖的训练策略。本文并不宣称复现上述模型论文，而是在相同 PCB 数据集和可比训练设置下进行应用层实验，观察不同输入分辨率和模型版本在工业小目标场景中的实际表现。"
    )
    add_heading(doc, "1.3 本文主要工作", 2)
    add_bullets(
        doc,
        [
            "完成 PCB 缺陷数据集从 COCO 标注到 YOLO txt 标注的转换，并按 70/20/10 划分训练集、验证集和测试集。",
            "编写并运行数据集审计脚本，统计各划分样本数、目标框数量、类别分布和目标尺度分布，补充报告中的数据分析工作量。",
            "完成 YOLO11n 在 640 输入下的 20 epoch 试验与 100 epoch 正式 baseline，分析训练轮数对收敛和指标的影响。",
            "完成 YOLO11n 在 960 输入下的服务器训练实验，验证输入分辨率提升对 PCB 小目标检测的影响。",
            "完成 YOLO26n-960 对比实验，讨论通用模型升级在特定工业数据集上可能不稳定的原因。",
            "基于混淆矩阵、预测图和指标差异进行误差分析，并提出自适应切片、困难样本挖掘、轻量化部署评估等后续扩展方向。",
        ],
    )

    add_heading(doc, "2 计算机视觉基础理论与相关算法", 1)
    add_heading(doc, "2.1 目标检测任务定义", 2)
    add_body_paragraph(
        doc,
        "目标检测要求模型同时预测目标类别和目标位置。对于 PCB 缺陷检测而言，输入为 PCB 图像，输出为若干带类别标签和置信度的边界框。常用评价指标包括 Precision、Recall、AP、mAP@50 和 mAP@50:95。其中 Precision 反映预测为缺陷的框中有多少是真缺陷，Recall 反映真实缺陷中有多少被检出；mAP@50 使用 IoU 阈值 0.5 评估检测效果，mAP@50:95 则在多个 IoU 阈值上取平均，更能体现定位质量。"
    )
    add_heading(doc, "2.2 YOLO 系列检测模型", 2)
    add_body_paragraph(
        doc,
        "YOLO 系列模型将目标检测建模为单阶段回归与分类问题，通常由 Backbone、Neck 和 Head 三部分组成。Backbone 负责提取图像特征，Neck 通过特征金字塔或路径聚合结构融合不同尺度信息，Head 输出边界框、类别概率和置信度。PCB 缺陷检测需要依赖高分辨率局部纹理，多尺度特征融合和小目标正样本分配对召回率具有重要影响。"
    )
    add_heading(doc, "2.3 小目标检测难点", 2)
    add_body_paragraph(
        doc,
        "小目标检测的核心困难在于目标像素占比低、特征层下采样后信息容易丢失、标注框边界误差对 IoU 影响更大、背景纹理与目标纹理相似。PCB 图像中线路、焊盘和铜箔纹理高度重复，缺陷往往表现为局部断裂、缺口或多余铜皮，模型既要识别细粒度形态，又要避免把正常线路结构误判为缺陷。因此，本报告将分辨率提升、模型对比和误差分析作为扩展实验主线。"
    )

    add_heading(doc, "3 数据集构建与实验方案", 1)
    add_heading(doc, "3.1 数据集来源与类别", 2)
    add_body_paragraph(
        doc,
        f"本实验使用的 PCB 缺陷数据在本地整理为 YOLO 格式，共 {dataset['total_images']} 张图像、{dataset['total_boxes']} 个缺陷框，包含 missing_pad、mouse_bite、open_circuit、short、spur、spurious_copper 六类缺陷。按照随机种子 42 进行 70/20/10 划分，保证训练、验证和测试阶段具有相互独立的样本。"
    )
    split_rows = []
    for split in ["train", "val", "test"]:
        stats = dataset["splits"][split]
        split_rows.append([split, str(stats["images"]), str(stats["boxes"])])
    add_table(doc, ["划分", "图像数", "缺陷框数"], split_rows, "表 3-1 PCB 缺陷数据集划分统计")
    add_picture(doc, ASSET_DIR / "dataset_class_distribution.png", "图 3-1 各划分中六类缺陷框数量分布", width=6.1)
    add_picture(doc, ASSET_DIR / "box_area_distribution.png", "图 3-2 缺陷框归一化面积分布", width=6.1)
    add_body_paragraph(
        doc,
        "从类别分布看，训练集、验证集和测试集均覆盖六类缺陷，避免了某一类别只出现在训练集或只出现在测试集的极端情况。从目标框面积分布看，较多缺陷框属于小面积区域，这与后续实验中 mAP@50:95 提升有限、Recall 增益不明显的现象相互印证。"
    )

    add_heading(doc, "3.2 数据转换与质量检查", 2)
    add_body_paragraph(
        doc,
        "原始标注采用 COCO JSON 格式，转换脚本读取图像宽高和 COCO 边界框，将 x、y、w、h 转换为 YOLO 所需的类别编号、归一化中心点坐标和归一化宽高。转换完成后，使用数据审计脚本统计每个 split 的图像数、标签文件数、空标签数量、类别计数和总框数，并通过标签可视化确认标注框与缺陷区域对齐。"
    )
    add_picture(doc, ASSET_DIR / "method_pipeline.png", "图 3-3 本课程项目实验流程扩展图", width=6.2)

    add_heading(doc, "3.3 实验环境与训练设置", 2)
    env_rows = [
        ["本地 GPU", "NVIDIA GeForce RTX 3060 Laptop GPU 6GB", "YOLO11n-640 训练与调试"],
        ["服务器 GPU", "CUDA GPU 服务器", "YOLO11n-960 与 YOLO26n-960 训练"],
        ["深度学习框架", "PyTorch 2.11.0+cu128，Ultralytics 8.4.54", "模型训练、验证与结果导出"],
        ["数据格式", "YOLO txt + dataset yaml", "目标检测训练输入"],
        ["主要脚本", "convert / inspect / train / summarize / compare", "数据转换、训练和实验汇总"],
    ]
    add_table(doc, ["项目", "配置", "用途"], env_rows, "表 3-2 实验软硬件与代码环境")
    hyper_rows = [
        ["YOLO11n-640-e20", "yolo11n.pt", "640", "20", "4", "快速收敛试验"],
        ["YOLO11n-640", "yolo11n.pt", "640", "100", "4", "正式 baseline"],
        ["YOLO11n-960", "yolo11n.pt", "960", "100", "8", "高分辨率对照"],
        ["YOLO26n-960", "yolo26n.pt", "960", "100", "8", "模型版本对比"],
    ]
    add_table(doc, ["实验名称", "模型", "输入尺寸", "Epoch", "Batch", "实验目的"], hyper_rows, "表 3-3 训练超参数与实验目的")
    add_body_paragraph(
        doc,
        "四组实验均固定随机种子为 42，采用 Ultralytics 默认优化器与增强策略，保存训练曲线、混淆矩阵、预测可视化图和 best.pt 权重。这样可以在不引入过多额外变量的前提下，将报告扩展为清晰的控制变量实验。"
    )

    add_heading(doc, "4 实验结果与扩展分析", 1)
    add_heading(doc, "4.1 整体指标对比", 2)
    metric_rows = []
    for name, metrics in runs.items():
        final = metrics["final"]
        metric_rows.append(
            [
                name,
                str(metrics["epochs"]),
                fmt(final["precision"]),
                fmt(final["recall"]),
                fmt(final["map50"]),
                fmt(final["map5095"]),
                fmt(metrics["best_map50"]["value"]),
                fmt(metrics["best_map5095"]["value"]),
            ]
        )
    add_table(
        doc,
        ["实验", "Epoch", "Precision", "Recall", "mAP@50", "mAP@50:95", "最佳 mAP@50", "最佳 mAP@50:95"],
        metric_rows,
        "表 4-1 YOLO11n/YOLO26n PCB 缺陷检测实验结果对比",
    )
    add_picture(doc, ASSET_DIR / "metrics_comparison.png", "图 4-1 不同实验设置的验证指标对比", width=6.2)
    add_body_paragraph(
        doc,
        "20 epoch 的 YOLO11n-640 仅达到 0.46298 mAP@50，说明该任务在短训练轮数下尚未充分收敛。训练到 100 epoch 后，mAP@50 提升到 0.79002，mAP@50:95 提升到 0.41293，表明延长训练轮数是必要的基础工作。进一步将输入尺寸从 640 提升到 960 后，YOLO11n 的 Precision、Recall、mAP@50 和 mAP@50:95 均有提升，其中 Precision 增幅最明显，说明更高分辨率有助于模型减少误检并捕获更多局部细节。"
    )
    add_body_paragraph(
        doc,
        "但需要注意的是，YOLO11n-960 的 Recall 仅比 YOLO11n-640 提高 0.03326，mAP@50:95 仅提高 0.02741。这说明提高全图输入尺寸虽然有效，但仍受到显存、训练成本和特征下采样的限制；对于 PCB 这类局部小缺陷任务，后续更值得探索切片检测、多尺度训练或小目标检测头等针对性方法。"
    )

    add_heading(doc, "4.2 YOLO11n 与 YOLO26n 的应用对比", 2)
    add_body_paragraph(
        doc,
        "YOLO26n 在官方设计目标上强调端到端、轻量检测头和小目标正样本覆盖，但在本实验的 PCB 小数据集、100 epoch、960 输入设置下，最终 mAP@50 为 0.76921，低于 YOLO11n-960 的 0.85118。该结果说明模型的通用 COCO 性能并不能直接等价为工业小目标数据集上的性能，模型版本更新也需要结合数据规模、训练策略、类别形态和评价指标重新验证。"
    )
    add_body_paragraph(
        doc,
        "从课程汇报角度看，这一结果并不是失败，而是很好的分析材料：它提示后续实验不能只追求更“新”的模型名称，而要围绕 PCB 缺陷的真实痛点设计实验，例如提高小目标召回率、改善边界框定位、降低漏检和误检成本。报告因此形成了更完整的研究闭环。"
    )

    add_heading(doc, "4.3 训练曲线与定性结果", 2)
    add_picture(doc, PROJECT_ROOT / "runs" / "baselines" / "yolo11n_full_960_server" / "results.png", "图 4-2 YOLO11n-960 训练曲线与验证指标变化", width=6.1)
    add_picture(doc, PROJECT_ROOT / "runs" / "baselines" / "yolo11n_full_960_server" / "confusion_matrix_normalized.png", "图 4-3 YOLO11n-960 归一化混淆矩阵", width=5.4)
    add_picture(doc, PROJECT_ROOT / "runs" / "baselines" / "yolo11n_full_960_server" / "val_batch0_pred.jpg", "图 4-4 YOLO11n-960 验证集预测可视化样例", width=6.1)
    add_body_paragraph(
        doc,
        "训练曲线显示，YOLO11n-960 的指标在中后期仍存在波动，最佳 mAP@50 出现在第 73 epoch，最佳 mAP@50:95 出现在第 86 epoch，说明模型在高分辨率条件下仍有继续调参空间。混淆矩阵和预测图可用于分析具体类别的漏检和误检：如果某类缺陷经常被划为 background，说明该类召回不足；如果类别之间互相混淆，则需要增加类别判别特征或进行困难样本增强。"
    )

    add_heading(doc, "4.4 扩展实验小结", 2)
    add_bullets(
        doc,
        [
            "训练轮数扩展：20 epoch 只能证明流程可跑通，100 epoch 才能形成可汇报的正式 baseline。",
            "输入分辨率扩展：960 输入明显优于 640 输入，但召回率提升有限，说明单纯放大全图不是最终答案。",
            "模型版本扩展：YOLO26n-960 在本数据集上未超过 YOLO11n-960，提示需要任务驱动的模型选择，而非简单追新。",
            "分析维度扩展：补充类别分布、目标尺度、训练曲线、混淆矩阵和预测图，使课程报告具备完整实验证据链。",
        ],
    )

    add_heading(doc, "5 误差分析与后续改进方案", 1)
    add_heading(doc, "5.1 误差来源分析", 2)
    add_body_paragraph(
        doc,
        "结合指标差异和可视化结果，当前模型主要存在三类问题。第一，小缺陷在全图缩放后细节不足，导致 Recall 增长慢；第二，PCB 线路纹理重复，模型可能把正常结构误判为缺陷，也可能把局部缺陷并入背景；第三，部分类别之间外观相似，例如短路、残铜和毛刺都可能表现为局部多余铜皮，类别边界并非完全清晰。"
    )
    add_heading(doc, "5.2 自适应切片检测方案", 2)
    add_body_paragraph(
        doc,
        "后续可将完整图像切分为若干重叠局部窗口，在局部窗口上训练或推理，再将预测框映射回原图并进行去重融合。固定切片能够提升小目标在输入图像中的相对像素占比，但会增加推理次数；自适应切片则可以先利用低分辨率模型筛选疑似区域，再对高风险局部区域进行高分辨率二次检测。其研究假设是：全图检测负责快速覆盖，局部检测负责提升小缺陷召回率和定位质量。"
    )
    add_heading(doc, "5.3 困难样本挖掘与类别均衡", 2)
    add_body_paragraph(
        doc,
        "针对被误判为 background 或类别混淆严重的样本，可以建立错误样本池，按类别和误差类型重新采样训练。对容易混淆的类别，可尝试更强的数据增强、局部裁剪增强、类别重加权损失或 focal-style 思路，使模型更关注困难样本。对样本量有限的课程项目而言，整理失败样例并进行定向增强，比盲目更换大模型更有解释性。"
    )
    add_heading(doc, "5.4 部署与工程评估", 2)
    add_body_paragraph(
        doc,
        "除精度指标外，工业检测还需要关注推理延迟、显存占用、模型大小和部署格式。后续可导出 ONNX 或 TensorRT，比较 YOLO11n-960、YOLO26n-960 和切片策略在单张图像上的平均推理时间。若切片策略提高召回但显著增加延迟，则需要进一步设计候选区域筛选或轻量化剪枝，以满足生产线实时检测需求。"
    )

    add_heading(doc, "6 总结与展望", 1)
    add_body_paragraph(
        doc,
        "本文围绕 PCB 缺陷小目标检测完成了一个可复现的计算机视觉课程项目。相比最初只完成 YOLO11n baseline 的工作，扩展后项目包含数据集转换、数据审计、标签质量检查、训练轮数对比、输入分辨率消融、YOLO11n 与 YOLO26n 应用对比、训练曲线与混淆矩阵分析、预测可视化和后续自适应切片方案设计。实验表明，YOLO11n-960 是当前设置下表现最好的模型，达到 0.85118 mAP@50 和 0.44034 mAP@50:95；输入分辨率提升有效，但小目标召回和定位质量仍有提升空间。"
    )
    add_body_paragraph(
        doc,
        "未来工作可从三个方向继续推进：一是实现固定切片和自适应切片推理，与全图 960 baseline 进行严格对比；二是针对漏检类别开展困难样本挖掘和局部增强；三是加入推理速度、模型大小和部署格式评估，使课程项目从算法实验进一步延伸到工业应用原型。"
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "[1] Gonzalez R C, Woods R E. Digital Image Processing[M]. 4th ed. New York: Pearson, 2018.",
        "[2] Szeliski R. Computer Vision: Algorithms and Applications[M]. 2nd ed. Cham: Springer, 2022.",
        "[3] Lowe D G. Distinctive image features from scale-invariant keypoints[J]. International Journal of Computer Vision, 2004, 60(2): 91-110.",
        "[4] Dalal N, Triggs B. Histograms of oriented gradients for human detection[C]//CVPR. 2005: 886-893.",
        "[5] Girshick R, Donahue J, Darrell T, et al. Rich feature hierarchies for accurate object detection and semantic segmentation[C]//CVPR. 2014.",
        "[6] Girshick R. Fast R-CNN[C]//ICCV. 2015.",
        "[7] Ren S, He K, Girshick R, et al. Faster R-CNN: Towards real-time object detection with region proposal networks[J]. IEEE TPAMI, 2017, 39(6): 1137-1149.",
        "[8] He K, Gkioxari G, Dollár P, et al. Mask R-CNN[C]//ICCV. 2017.",
        "[9] Liu W, Anguelov D, Erhan D, et al. SSD: Single Shot MultiBox Detector[C]//ECCV. 2016.",
        "[10] Lin T Y, Goyal P, Girshick R, et al. Focal loss for dense object detection[C]//ICCV. 2017.",
        "[11] Redmon J, Divvala S, Girshick R, et al. You Only Look Once: Unified, real-time object detection[C]//CVPR. 2016.",
        "[12] Redmon J, Farhadi A. YOLO9000: Better, faster, stronger[C]//CVPR. 2017.",
        "[13] Redmon J, Farhadi A. YOLOv3: An incremental improvement[EB/OL]. arXiv:1804.02767, 2018.",
        "[14] Bochkovskiy A, Wang C Y, Liao H Y M. YOLOv4: Optimal speed and accuracy of object detection[EB/OL]. arXiv:2004.10934, 2020.",
        "[15] Wang C Y, Bochkovskiy A, Liao H Y M. YOLOv7: Trainable bag-of-freebies sets new state-of-the-art for real-time object detectors[C]//CVPR. 2023.",
        "[16] Ge Z, Liu S, Wang F, et al. YOLOX: Exceeding YOLO series in 2021[EB/OL]. arXiv:2107.08430, 2021.",
        "[17] Jocher G, Qiu J. Ultralytics YOLO11[EB/OL]. 2024. https://github.com/ultralytics/ultralytics.",
        "[18] Jocher G, Qiu J, Liu M, et al. Ultralytics YOLO26: Unified real-time end-to-end vision models[EB/OL]. arXiv:2606.03748, 2026.",
        "[19] Zhao Y, Lv W, Xu S, et al. DETRs Beat YOLOs on real-time object detection[EB/OL]. arXiv:2304.08069, 2023.",
        "[20] Carion N, Massa F, Synnaeve G, et al. End-to-end object detection with transformers[C]//ECCV. 2020.",
        "[21] Lin T Y, Dollár P, Girshick R, et al. Feature pyramid networks for object detection[C]//CVPR. 2017.",
        "[22] He K, Zhang X, Ren S, et al. Deep residual learning for image recognition[C]//CVPR. 2016.",
        "[23] Huang W, Wei P. A PCB Dataset for Defects Detection and Classification[EB/OL]. arXiv:1901.08204, 2019.",
        "[24] Everingham M, Van Gool L, Williams C K I, et al. The PASCAL Visual Object Classes challenge[J]. IJCV, 2010, 88(2): 303-338.",
        "[25] Lin T Y, Maire M, Belongie S, et al. Microsoft COCO: Common objects in context[C]//ECCV. 2014.",
        "[26] Loshchilov I, Hutter F. Decoupled weight decay regularization[C]//ICLR. 2019.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(-0.74)
        p.paragraph_format.left_indent = Cm(0.74)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        set_run_font(run, size=10.5, name="宋体")

    add_heading(doc, "附录 A 核心代码与汇报口径", 1)
    add_body_paragraph(doc, "本项目关键脚本包括：scripts/convert_pcb_defect_coco_to_yolo.py、scripts/inspect_yolo_dataset.py、scripts/train_yolo_baseline.py、scripts/summarize_yolo_run.py、scripts/compare_yolo_runs.py、scripts/build_report_assets.py。汇报时建议围绕“为什么需要扩展、做了哪些对照、结果说明了什么、后续如何改进”四个问题展开。")
    add_body_paragraph(doc, "推荐汇报结论：本项目不是简单调用 YOLO 训练一次模型，而是完成了 PCB 缺陷检测从数据准备、模型训练到结果诊断的闭环。当前最优设置为 YOLO11n-960，后续研究重点应从全图输入转向局部切片和小目标召回率提升。")

    doc.save(OUTPUT_PATH)
    print(f"Wrote report: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_report()
